# -*- coding: utf-8 -*-
"""Парсинг корпуса «Источники информации» → documents.jsonl.

Один JSON на строку: {doc_id, source_path, source_type, title, authors,
year, lang, text, chunks, meta}. Формат — см. BRIEF_DATA_Dubinin.md и PIPELINE.md.

Принципы:
- один битый файл не роняет прогон: каждый файл в try/except, ошибки в лог;
- сканы (PDF без текстового слоя) пропускаются и логируются, OCR не делаем;
- дедупликация по хешу нормализованного текста;
- метаданные (год, журнал, выпуск) — из пути и имени файла.

Запуск:
    python -m ingest.parse --input data/raw --output data/processed/documents.jsonl
    python -m ingest.parse --priority-only          # только Обзоры + Статьи + Журналы
    python -m ingest.parse --limit 20               # быстрый прогон
"""
import argparse
import hashlib
import json
import logging
import re
import sys
import zipfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from langdetect import DetectorFactory, detect

DetectorFactory.seed = 0  # детерминированное определение языка

log = logging.getLogger("ingest")

# ---------------------------------------------------------------- настройки

CHUNK_WORDS = 700      # целевой размер куска, слов
CHUNK_OVERLAP = 100    # перекрытие между кусками, слов
MIN_TEXT_CHARS = 200   # меньше — считаем документ пустым (скан/мусор)

# верхняя папка корпуса → source_type
SOURCE_TYPES = {
    "журналы": "journal",
    "материалы конференций": "conference",
    "доклады": "report",
    "обзоры": "review",
    "статьи": "article",
}

# приоритет обработки: Обзоры (готовые литобзоры) → Статьи → Журналы → остальное
PRIORITY = {"review": 0, "article": 1, "journal": 2, "report": 3, "conference": 4}

SKIP_EXTS = {
    ".xls", ".xlsx", ".xlsm",          # рыночная статистика — отдельный трек
    ".pptx", ".ppt", ".pptm", ".docm",  # презентации/макросы — не текстовый корпус
    ".gif", ".jpg", ".png",
    ".001", ".002",                    # куски многотомных архивов
    ".rar",                            # нужен unrar; логируем и идём дальше
    ".doc",                            # старый формат; их мало — пропускаем по ТЗ
}

TRANSLIT = str.maketrans({
    "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e", "ё": "e",
    "ж": "zh", "з": "z", "и": "i", "й": "i", "к": "k", "л": "l", "м": "m",
    "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t", "у": "u",
    "ф": "f", "х": "h", "ц": "c", "ч": "ch", "ш": "sh", "щ": "sch",
    "ъ": "", "ы": "y", "ь": "", "э": "e", "ю": "yu", "я": "ya",
})


# ---------------------------------------------------------------- извлечение текста

def parse_pdf(path: Path) -> str:
    """Текст из PDF постранично. Пустой результат = скан (текстового слоя нет)."""
    parts = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def parse_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    # текст из таблиц тоже ценен (составы, параметры)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def docx_meta(path: Path) -> dict:
    """Заголовок из свойств docx, если заполнен.

    Поле author намеренно не берём: там создатель файла (один и тот же
    сотрудник на весь корпус обзоров), а не авторы работы.
    """
    out = {}
    try:
        props = Document(str(path)).core_properties
        if props.title:
            out["title"] = props.title.strip()
    except Exception:
        pass
    return out


# ---------------------------------------------------------------- метаданные

def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ").replace("­", "")  # nbsp, мягкий перенос
    text = re.sub(r"-\n(?=[а-яa-z])", "", text)             # перенос слова со строки
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def detect_lang(text: str) -> str:
    sample = text[:2000]
    try:
        lang = detect(sample)
    except Exception:
        return "ru"
    return lang if lang in ("ru", "en") else ("ru" if re.search(r"[а-яА-Я]", sample) else "en")


def year_from(rel_path: str, text: str) -> int | None:
    """Год: сперва из пути/имени файла, затем из начала текста."""
    for m in re.finditer(r"(?<!\d)(19[89]\d|20[0-4]\d)(?!\d)", rel_path):
        return int(m.group(1))
    # короткий год в имени журнала: "№ 03_24.pdf"
    m = re.search(r"№\s*\d+[_-](\d{2})\b", rel_path)
    if m:
        return 2000 + int(m.group(1))
    m = re.search(r"(?<!\d)(19[89]\d|20[0-4]\d)(?!\d)", text[:3000])
    return int(m.group(1)) if m else None


def path_meta(rel_path: str) -> dict:
    """Журнал и выпуск из пути вида 'Журналы/Горная промышленность/2024/ГП № 3-2024.pdf'."""
    meta = {}
    parts = rel_path.split("/")
    if parts[0].lower() == "журналы" and len(parts) > 2:
        meta["journal"] = parts[1]
    if parts[0].lower() == "материалы конференций" and len(parts) > 2:
        meta["conference"] = parts[1]
    m = re.search(r"№\s*(\d+)[_-](\d{2,4})", parts[-1])
    if m:
        yy = m.group(2)
        meta["issue"] = f"{int(m.group(1))}-{yy if len(yy) == 4 else '20' + yy}"
    return meta


def guess_title(text: str, path: Path) -> str:
    """Имя файла как заголовок: в этом корпусе оно описательное
    («Методы очистки шахтных вод»), а первые строки текста — лист согласования."""
    stem = re.sub(r"[_]+", " ", path.stem).strip()
    return stem or path.stem


def make_doc_id(rel_path: str) -> str:
    """Читаемый стабильный id: транслит имени + короткий хеш пути."""
    stem = Path(rel_path).stem.lower().translate(TRANSLIT)
    slug = re.sub(r"[^a-z0-9]+", "_", stem).strip("_")[:60]
    h = hashlib.sha1(rel_path.encode("utf-8")).hexdigest()[:8]
    return f"{slug}_{h}" if slug else h


# ---------------------------------------------------------------- нарезка на чанки

def chunk_text(text: str, size: int = CHUNK_WORDS, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Куски ~size слов с перекрытием ~overlap, по границам абзацев."""
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks, current, count = [], [], 0
    for para in paragraphs:
        words = para.split()
        # огромный абзац режем жёстко по словам
        while len(words) > size:
            if current:
                chunks.append(" ".join(current))
                current, count = [], 0
            chunks.append(" ".join(words[:size]))
            words = words[size - overlap:]
        current.append(" ".join(words))
        count += len(words)
        if count >= size:
            chunks.append("\n".join(current))
            tail = " ".join(current).split()[-overlap:]
            current, count = [" ".join(tail)], len(tail)
    if current and count > 20:
        chunks.append("\n".join(current))
    elif current and chunks:
        chunks[-1] += "\n" + "\n".join(current)
    return chunks or ([text] if text else [])


# ---------------------------------------------------------------- сборка записи

def make_record(path: Path, root: Path) -> dict | None:
    """Одна запись documents.jsonl. None = файл пропущен (причина в логе)."""
    rel = path.relative_to(root).as_posix()
    ext = path.suffix.lower()

    if ext == ".pdf":
        text = parse_pdf(path)
        extra = {}
    elif ext == ".docx":
        text = parse_docx(path)
        extra = docx_meta(path)
    else:
        log.info("skip (unsupported %s): %s", ext, rel)
        return None

    text = clean_text(text)
    if len(text) < MIN_TEXT_CHARS:
        log.warning("skip (no text layer / empty, %d chars): %s", len(text), rel)
        return None

    top = rel.split("/")[0].lower()
    source_type = SOURCE_TYPES.get(top, "reference")

    meta = path_meta(rel)
    # у выпусков журналов имя файла криптичное («№ 01_24») — собираем из метаданных
    if meta.get("journal") and meta.get("issue"):
        journal_title = f"{meta['journal']} № {meta['issue']}"
    else:
        journal_title = None

    record = {
        "doc_id": make_doc_id(rel),
        "source_path": rel,
        "source_type": source_type,
        "title": extra.get("title") or journal_title or guess_title(text, path),
        "authors": extra.get("authors", []),
        "year": year_from(rel, text),
        "lang": detect_lang(text),
        "text": text,
        "chunks": chunk_text(text),
        "meta": meta,
    }
    return record


# ---------------------------------------------------------------- вложенные архивы

def unpack_nested_zips(root: Path) -> None:
    """Zip внутри корпуса разворачиваем в папку рядом (один раз), rar только логируем."""
    for zp in list(root.rglob("*.zip")):
        target = zp.with_suffix("")
        if target.exists():
            continue
        try:
            with zipfile.ZipFile(zp) as z:
                for info in z.infolist():
                    if info.is_dir():
                        continue
                    name = info.filename
                    try:
                        name = name.encode("cp437").decode("cp866")
                    except (UnicodeEncodeError, UnicodeDecodeError):
                        pass
                    out = target / name
                    out.parent.mkdir(parents=True, exist_ok=True)
                    with z.open(info) as src, open(out, "wb") as dst:
                        dst.write(src.read())
            log.info("unpacked zip: %s", zp.relative_to(root).as_posix())
        except Exception as e:
            log.error("bad zip %s: %s", zp.relative_to(root).as_posix(), e)
    for rp in root.rglob("*.rar"):
        log.info("skip rar (unrar not available): %s", rp.relative_to(root).as_posix())


# ---------------------------------------------------------------- прогон

def collect_files(root: Path, priority_only: bool) -> list[Path]:
    files = []
    for p in root.rglob("*"):
        if not p.is_file() or p.suffix.lower() not in (".pdf", ".docx"):
            continue
        if p.name.startswith("~$"):  # временные файлы Word
            continue
        top = p.relative_to(root).as_posix().split("/")[0].lower()
        stype = SOURCE_TYPES.get(top, "reference")
        if priority_only and stype not in ("review", "article", "journal"):
            continue
        files.append(p)
    files.sort(key=lambda p: (
        PRIORITY.get(SOURCE_TYPES.get(p.relative_to(root).as_posix().split("/")[0].lower(), ""), 9),
        p.suffix.lower() != ".docx",  # docx раньше pdf внутри одной категории
        p.as_posix(),
    ))
    return files


def normalized_hash(text: str) -> str:
    norm = re.sub(r"\s+", " ", text.lower())[:20000]
    return hashlib.sha1(norm.encode("utf-8")).hexdigest()


def main():
    ap = argparse.ArgumentParser(description="Парсинг корпуса в documents.jsonl")
    ap.add_argument("--input", default="data/raw")
    ap.add_argument("--output", default="data/processed/documents.jsonl")
    ap.add_argument("--limit", type=int, default=0, help="обработать не больше N файлов")
    ap.add_argument("--priority-only", action="store_true",
                    help="только Обзоры + Статьи + Журналы")
    ap.add_argument("--no-unpack", action="store_true", help="не разворачивать вложенные zip")
    args = ap.parse_args()

    root = Path(args.input)
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    logging.basicConfig(
        level=logging.INFO,
        format="%(levelname)s %(message)s",
        handlers=[
            logging.StreamHandler(sys.stderr),
            logging.FileHandler(out_path.parent / "parse_errors.log", encoding="utf-8"),
        ],
    )

    if not args.no_unpack:
        unpack_nested_zips(root)

    files = collect_files(root, args.priority_only)
    if args.limit:
        files = files[:args.limit]
    log.info("files to process: %d", len(files))

    seen_hashes: set[str] = set()
    written = errors = dupes = skipped = 0
    with open(out_path, "w", encoding="utf-8") as f:
        for i, path in enumerate(files, 1):
            try:
                record = make_record(path, root)
            except Exception as e:
                errors += 1
                log.error("FAILED %s: %r", path.relative_to(root).as_posix(), e)
                continue
            if record is None:
                skipped += 1
                continue
            h = normalized_hash(record["text"])
            if h in seen_hashes:
                dupes += 1
                log.info("skip duplicate: %s", record["source_path"])
                continue
            seen_hashes.add(h)
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
            written += 1
            if i % 25 == 0:
                log.info("progress: %d/%d (written %d)", i, len(files), written)

    log.info("DONE: written=%d skipped=%d duplicates=%d errors=%d -> %s",
             written, skipped, dupes, errors, out_path)


if __name__ == "__main__":
    main()
